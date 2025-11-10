#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Génère un document Word à partir du fichier Markdown Workflow_Natclinn.md
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, url, text):
    """Ajoute un hyperlien à un paragraphe"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def parse_markdown_to_word(md_file, docx_file):
    """Convertit le fichier Markdown en document Word"""
    
    # Lire le fichier Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Créer le document Word
    doc = Document()
    
    # Définir les styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Parser le Markdown ligne par ligne
    lines = content.split('\n')
    in_code_block = False
    in_table = False
    table_obj = None
    table_headers = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Bloc de code
        if line.startswith('```'):
            in_code_block = not in_code_block
            if not in_code_block:
                doc.add_paragraph()  # Espace après le bloc de code
            i += 1
            continue
        
        if in_code_block:
            p = doc.add_paragraph(line, style='Normal')
            p.paragraph_format.left_indent = Inches(0.5)
            if p.runs:
                run = p.runs[0]
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 0, 139)
            i += 1
            continue
        
        # Titre niveau 1
        if line.startswith('# ') and not line.startswith('##'):
            text = line[2:].strip()
            p = doc.add_heading(text, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Titre niveau 2
        if line.startswith('## '):
            text = line[3:].strip()
            doc.add_heading(text, level=1)
            i += 1
            continue
        
        # Titre niveau 3
        if line.startswith('### '):
            text = line[4:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue
        
        # Séparateur horizontal
        if line.strip() == '---':
            doc.add_paragraph()
            i += 1
            continue
        
        # Tableau
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                # Début du tableau - extraire les en-têtes
                table_headers = [cell.strip() for cell in line.split('|')[1:-1]]
                in_table = True
                # Ignorer la ligne de séparation
                i += 2
                # Créer le tableau
                table_obj = doc.add_table(rows=1, cols=len(table_headers))
                table_obj.style = 'Light Grid Accent 1'
                # Remplir les en-têtes
                for idx, header in enumerate(table_headers):
                    cell = table_obj.rows[0].cells[idx]
                    cell.text = header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                continue
            else:
                # Ligne de données du tableau
                cells_data = [cell.strip() for cell in line.split('|')[1:-1]]
                row = table_obj.add_row()
                for idx, cell_text in enumerate(cells_data):
                    row.cells[idx].text = cell_text
                i += 1
                continue
        else:
            if in_table:
                in_table = False
                doc.add_paragraph()  # Espace après le tableau
        
        # Liste à puces
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Détecter les liens
            link_match = re.search(r'\*\*(.*?)\*\* : (.*)', text)
            if link_match:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(link_match.group(1)).bold = True
                p.add_run(' : ' + link_match.group(2))
            else:
                doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        
        # Texte gras avec **
        if '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            i += 1
            continue
        
        # Texte code inline avec `
        if '`' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(`[^`]+`)', line)
            for part in parts:
                if part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(220, 20, 60)
                else:
                    p.add_run(part)
            i += 1
            continue
        
        # Ligne vide
        if line.strip() == '':
            i += 1
            continue
        
        # Paragraphe normal
        doc.add_paragraph(line.strip())
        i += 1
    
    # Sauvegarder le document
    doc.save(docx_file)
    print(f"Document Word généré : {docx_file}")

if __name__ == '__main__':
    md_file = r'C:\var\www\natclinn\results\Workflow_Natclinn.md'
    docx_file = r'C:\var\www\natclinn\results\Workflow_Natclinn.docx'
    
    try:
        parse_markdown_to_word(md_file, docx_file)
        print("✓ Conversion réussie !")
    except Exception as e:
        print(f"✗ Erreur : {e}")
        import traceback
        traceback.print_exc()
